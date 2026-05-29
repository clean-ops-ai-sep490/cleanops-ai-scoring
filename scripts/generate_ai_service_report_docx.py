from __future__ import annotations

import json
from pathlib import Path
from typing import Iterable

import matplotlib

matplotlib.use("Agg")

import matplotlib.pyplot as plt
from matplotlib.patches import FancyArrowPatch, FancyBboxPatch

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
REPORT_DIR = ROOT / "docs" / "report"
GENERATED_DIR = REPORT_DIR / "generated"
OUT_DOCX = REPORT_DIR / "cleanops_ai_service_report_en.docx"
UNET_BENCHMARK_SOURCE = Path(r"E:\capstone\train_train_train\cleanops-ai\docs\benchmark_report_vi.md")

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
MUTED = "667085"
LIGHT_GRAY = "F2F4F7"
PALE_BLUE = "E8EEF5"
PALE_GOLD = "FFF4D6"
PALE_RED = "FDECEC"
WHITE = "FFFFFF"

UNET_BENCHMARK = {
    "dataset": "data/processed/benchmark",
    "evaluated": "46/46",
    "missing_masks": "0",
    "checkpoint": "outputs/unet/runs/unet_resnet34_20260518_200957/best.pt",
    "selection_rule": "minimum best_val_loss parsed from training logs",
    "best_val_loss": 0.6757,
    "old": {
        "Pixel Accuracy": 0.947288,
        "Mean IoU": 0.372956,
        "Mean Dice/F1": 0.422034,
    },
    "new": {
        "Pixel Accuracy": 0.960080,
        "Mean IoU": 0.418637,
        "Mean Dice/F1": 0.478676,
    },
    "classes": [
        {
            "name": "background",
            "support": 604_257_149,
            "share": 95.28,
            "iou": 0.960370,
            "dice": 0.979785,
        },
        {
            "name": "dirty_area",
            "support": 15_416_922,
            "share": 2.43,
            "iou": 0.295541,
            "dice": 0.456243,
        },
        {
            "name": "wet_surface",
            "support": 14_487_258,
            "share": 2.28,
            "iou": 0.000000,
            "dice": 0.000000,
        },
    ],
    "total_pixels": 634_161_329,
    "confusion": [
        ["background", 604_141_410, 115_739, 0],
        ["dirty_area", 10_712_425, 4_704_497, 0],
        ["wet_surface", 14_101_657, 385_601, 0],
    ],
}


def read_json(path: Path) -> dict:
    return json.loads(path.read_text(encoding="utf-8-sig"))


def pct(value: float) -> str:
    return f"{float(value) * 100:.2f}%"


def one(value: float) -> str:
    return f"{float(value):.1f}"


def two(value: float) -> str:
    return f"{float(value):.2f}"


def set_run_font(run, *, size: float | None = None, color: str | None = None, bold: bool | None = None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold


def shade_cell(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="D0D5DD", size="6"):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def style_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    title = styles["Title"]
    title.font.name = "Calibri"
    title.font.size = Pt(24)
    title.font.bold = True
    title.font.color.rgb = RGBColor.from_string(INK)
    title.paragraph_format.space_after = Pt(8)

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def add_footer(doc: Document):
    section = doc.sections[0]
    header = section.header.paragraphs[0]
    header.text = "CleanOps AI Scoring Service"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run_font(header.runs[0], size=9, color=MUTED)

    footer = section.footer.paragraphs[0]
    footer.text = "Prepared for capstone defense | English technical report"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(footer.runs[0], size=9, color=MUTED)


def add_title_block(doc: Document):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run("TECHNICAL REPORT")
    set_run_font(r, size=10, color=MUTED, bold=True)

    title = doc.add_paragraph(style="Title")
    title.add_run("CleanOps AI Scoring Service")

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(16)
    r = subtitle.add_run(
        "Hybrid cleanliness verification with YOLO, domain-specific U-Net, "
        "and SAM3/Roboflow-style auxiliary segmentation"
    )
    set_run_font(r, size=13, color="344054")

    rows = [
        ("Document purpose", "Capstone defense report and implementation evidence"),
        ("Scope", "Cleanliness AI scoring service, benchmark evidence, retrain loop, and safety limitations"),
        ("Report date", "May 28, 2026"),
        ("Primary artifact", "docs/report/cleanops_ai_service_report_en.docx"),
    ]
    table = doc.add_table(rows=0, cols=2)
    table.autofit = False
    table.columns[0].width = Inches(1.7)
    table.columns[1].width = Inches(4.7)
    set_table_borders(table, color="E4E7EC")
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value
        shade_cell(cells[0], LIGHT_GRAY)
        for cell in cells:
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for p in cell.paragraphs:
                for run in p.runs:
                    set_run_font(run, size=10.5, color=INK if cell is cells[0] else "101828", bold=cell is cells[0])

    doc.add_paragraph()
    add_callout(
        doc,
        "Core thesis",
        "The service is intentionally hybrid: the foundation-style SAM3/Roboflow provider supplies broad auxiliary "
        "segmentation, while the U-Net trained for CleanOps captures domain-specific dirty and wet-floor patterns. "
        "The scoring layer unions evidence and routes uncertain broad-mask cases to supervisor review instead of "
        "blindly approving or rejecting them.",
        fill=PALE_BLUE,
    )


def add_callout(doc: Document, title: str, body: str, fill: str = PALE_BLUE):
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    table.columns[0].width = Inches(6.3)
    set_table_borders(table, color="D0D5DD")
    cell = table.cell(0, 0)
    shade_cell(cell, fill)
    set_cell_margins(cell, top=120, bottom=120, start=160, end=160)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    set_run_font(r, size=11, color=INK, bold=True)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(body)
    set_run_font(r2, size=10.3, color="101828")
    doc.add_paragraph()


def add_bullets(doc: Document, items: Iterable[str]):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(item)
        set_run_font(r, size=10.8, color="101828")


def add_caption(doc: Document, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    set_run_font(r, size=9.3, color=MUTED)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.autofit = False
    if widths:
        for idx, width in enumerate(widths):
            table.columns[idx].width = Inches(width)
    set_table_borders(table)
    header_cells = table.rows[0].cells
    set_repeat_table_header(table.rows[0])
    for idx, text in enumerate(headers):
        cell = header_cells[idx]
        cell.text = text
        shade_cell(cell, LIGHT_GRAY)
        set_cell_margins(cell)
        for p in cell.paragraphs:
            p.paragraph_format.space_after = Pt(0)
            for run in p.runs:
                set_run_font(run, size=9.5, color=INK, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for idx, text in enumerate(row):
            cells[idx].text = str(text)
            set_cell_margins(cells[idx])
            cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for p in cells[idx].paragraphs:
                p.paragraph_format.space_after = Pt(0)
                for run in p.runs:
                    set_run_font(run, size=9.2, color="101828")
    doc.add_paragraph()
    return table


def draw_box(ax, x, y, w, h, text, fill, edge="#2E74B5", size=10):
    patch = FancyBboxPatch(
        (x, y),
        w,
        h,
        boxstyle="round,pad=0.03,rounding_size=0.03",
        linewidth=1.2,
        edgecolor=edge,
        facecolor=fill,
    )
    ax.add_patch(patch)
    ax.text(x + w / 2, y + h / 2, text, ha="center", va="center", fontsize=size, color="#0B2545", wrap=True)


def arrow(ax, start, end):
    ax.add_patch(
        FancyArrowPatch(
            start,
            end,
            arrowstyle="-|>",
            mutation_scale=14,
            linewidth=1.2,
            color="#667085",
        )
    )


def create_system_diagram(path: Path):
    fig, ax = plt.subplots(figsize=(11, 5.7))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 5)
    ax.axis("off")
    draw_box(ax, 0.3, 3.2, 1.8, 0.85, "Cleaner uploads\ncompletion image", "#E8F3FF")
    draw_box(ax, 2.5, 3.2, 1.9, 0.85, "CleanOps Backend\njob orchestration", "#F2F4F7")
    draw_box(ax, 4.9, 3.2, 2.0, 0.85, "AI Scoring API\ninference + rules", "#E8EEF5")
    draw_box(ax, 7.4, 3.2, 2.0, 0.85, "Visualization Blob\nreview evidence", "#F6FEF9")
    draw_box(ax, 4.9, 1.8, 2.0, 0.85, "Model storage\nruntime cache\nprovider config", "#FFF4D6", size=9.0)
    draw_box(ax, 2.5, 0.6, 1.9, 0.85, "Supervisor review\nPASS/PENDING/FAIL", "#FDECEC")
    draw_box(ax, 4.9, 0.6, 2.0, 0.85, "Retrain bridge\napproved annotations", "#F2F4F7")
    draw_box(ax, 7.4, 0.6, 2.0, 0.85, "Promotion gate\ncandidate -> active", "#E8F3FF")
    arrow(ax, (2.1, 3.62), (2.5, 3.62))
    arrow(ax, (4.4, 3.62), (4.9, 3.62))
    arrow(ax, (6.9, 3.62), (7.4, 3.62))
    arrow(ax, (5.9, 3.2), (5.9, 2.65))
    arrow(ax, (5.0, 1.03), (4.4, 1.03))
    arrow(ax, (4.4, 1.03), (4.9, 1.03))
    arrow(ax, (6.9, 1.03), (7.4, 1.03))
    arrow(ax, (8.4, 1.45), (6.6, 3.2))
    fig.tight_layout()
    fig.savefig(path, dpi=220, bbox_inches="tight")
    plt.close(fig)


def create_fusion_diagram(path: Path):
    fig, ax = plt.subplots(figsize=(11, 5.8))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 5)
    ax.axis("off")
    draw_box(ax, 0.35, 2.05, 1.4, 0.8, "Input image", "#F2F4F7")
    draw_box(ax, 2.4, 3.5, 1.7, 0.75, "YOLO\ntrash/debris", "#F6FEF9")
    draw_box(ax, 2.4, 2.1, 1.7, 0.75, "U-Net\nspecialized dirty/wet", "#E8F3FF")
    draw_box(ax, 2.4, 0.9, 1.7, 0.75, "SAM3/Roboflow\nauxiliary masks", "#FFF4D6")
    draw_box(ax, 5.0, 2.1, 1.8, 0.85, "Union dirty/wet mask\n+ object penalty", "#E8EEF5")
    draw_box(ax, 7.6, 2.1, 1.8, 0.85, "Rule-based score\nPASS/PENDING/FAIL", "#FDECEC")
    arrow(ax, (1.75, 2.45), (2.4, 3.85))
    arrow(ax, (1.75, 2.45), (2.4, 2.48))
    arrow(ax, (1.75, 2.45), (2.4, 1.25))
    arrow(ax, (4.1, 3.85), (5.0, 2.75))
    arrow(ax, (4.1, 2.48), (5.0, 2.48))
    arrow(ax, (4.1, 1.25), (5.0, 2.2))
    arrow(ax, (6.8, 2.5), (7.6, 2.5))
    ax.text(
        5.9,
        0.45,
        "Calibration sends broad or uncertain masks to PENDING review\ninstead of over-claiming automatic PASS/FAIL.",
        ha="center",
        va="center",
        fontsize=9.5,
        color="#475467",
    )
    fig.tight_layout()
    fig.savefig(path, dpi=220, bbox_inches="tight")
    plt.close(fig)


def create_retrain_diagram(path: Path):
    fig, ax = plt.subplots(figsize=(11, 4.6))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 4)
    ax.axis("off")
    labels = [
        ("AI result\n+ visualization", "#E8EEF5"),
        ("Supervisor\nreview", "#FDECEC"),
        ("Approved\nannotations", "#FFF4D6"),
        ("Bridge dataset\nYOLO/U-Net", "#F2F4F7"),
        ("Candidate\ntraining", "#E8F3FF"),
        ("Benchmark +\npromotion gate", "#F6FEF9"),
    ]
    xs = [0.25, 1.9, 3.55, 5.2, 6.85, 8.5]
    for x, (label, fill) in zip(xs, labels):
        draw_box(ax, x, 1.8, 1.25, 0.85, label, fill, size=9.3)
    for x in xs[:-1]:
        arrow(ax, (x + 1.25, 2.23), (x + 1.65, 2.23))
    arrow(ax, (9.1, 1.8), (6.0, 0.75))
    ax.text(6.0, 0.5, "Only promoted candidates replace active artifacts; benchmark samples remain frozen.", ha="center", fontsize=9.2, color="#475467")
    fig.tight_layout()
    fig.savefig(path, dpi=220, bbox_inches="tight")
    plt.close(fig)


def create_metric_charts(path_rates: Path, path_latency: Path, active: dict, candidate: dict, roboflow: dict):
    variants = ["Active\nSAM3 off", "Candidate U-Net\nSAM3 off", "Roboflow merged\ncalibrated"]
    false_pass = [active["false_pass_rate"] * 100, candidate["false_pass_rate"] * 100, roboflow["false_pass_rate"] * 100]
    false_fail = [active["false_fail_rate"] * 100, candidate["false_fail_rate"] * 100, roboflow["false_fail_rate"] * 100]
    pending = [active["pending_review_rate"] * 100, candidate["pending_review_rate"] * 100, roboflow["pending_review_rate"] * 100]

    fig, ax = plt.subplots(figsize=(9, 4.8))
    x = range(len(variants))
    width = 0.24
    ax.bar([i - width for i in x], false_pass, width, label="False pass", color="#D92D20")
    ax.bar(list(x), false_fail, width, label="False fail", color="#F79009")
    ax.bar([i + width for i in x], pending, width, label="Pending review", color="#2E74B5")
    ax.set_xticks(list(x))
    ax.set_xticklabels(variants)
    ax.set_ylabel("Rate (%)")
    ax.set_ylim(0, 110)
    ax.set_title("Pilot benchmark safety and review rates")
    ax.grid(axis="y", alpha=0.25)
    ax.legend(loc="upper left", ncol=3)
    for bars in ax.containers:
        ax.bar_label(bars, fmt="%.1f", fontsize=8, padding=2)
    fig.tight_layout()
    fig.savefig(path_rates, dpi=220, bbox_inches="tight")
    plt.close(fig)

    latencies = [active["average_latency_ms"], candidate["average_latency_ms"], roboflow["average_latency_ms"]]
    fig, ax = plt.subplots(figsize=(8.5, 4.5))
    bars = ax.bar(variants, latencies, color=["#98A2B3", "#2E74B5", "#12B76A"])
    ax.set_ylabel("Average latency (ms)")
    ax.set_title("Average end-to-end API latency by variant")
    ax.grid(axis="y", alpha=0.25)
    ax.bar_label(bars, fmt="%.0f ms", fontsize=9, padding=3)
    fig.tight_layout()
    fig.savefig(path_latency, dpi=220, bbox_inches="tight")
    plt.close(fig)


def create_unet_benchmark_charts(path_improvement: Path, path_class_metrics: Path, path_support: Path):
    metrics = list(UNET_BENCHMARK["old"].keys())
    old_values = [UNET_BENCHMARK["old"][metric] for metric in metrics]
    new_values = [UNET_BENCHMARK["new"][metric] for metric in metrics]

    fig, ax = plt.subplots(figsize=(8.8, 4.6))
    x = range(len(metrics))
    width = 0.34
    old_bars = ax.bar([i - width / 2 for i in x], old_values, width, label="Previous benchmark", color="#98A2B3")
    new_bars = ax.bar([i + width / 2 for i in x], new_values, width, label="Best-val-loss checkpoint", color="#2E74B5")
    ax.set_xticks(list(x))
    ax.set_xticklabels(metrics)
    ax.set_ylim(0, 1.05)
    ax.set_ylabel("Score")
    ax.set_title("U-Net benchmark improvement on 46-image mask set")
    ax.grid(axis="y", alpha=0.25)
    ax.legend(loc="upper left")
    ax.bar_label(old_bars, fmt="%.3f", fontsize=8, padding=2)
    ax.bar_label(new_bars, fmt="%.3f", fontsize=8, padding=2)
    fig.tight_layout()
    fig.savefig(path_improvement, dpi=220, bbox_inches="tight")
    plt.close(fig)

    classes = [item["name"] for item in UNET_BENCHMARK["classes"]]
    ious = [item["iou"] for item in UNET_BENCHMARK["classes"]]
    dices = [item["dice"] for item in UNET_BENCHMARK["classes"]]
    fig, ax = plt.subplots(figsize=(8.8, 4.6))
    x = range(len(classes))
    iou_bars = ax.bar([i - width / 2 for i in x], ious, width, label="IoU", color="#12B76A")
    dice_bars = ax.bar([i + width / 2 for i in x], dices, width, label="Dice/F1", color="#2E74B5")
    ax.set_xticks(list(x))
    ax.set_xticklabels(classes)
    ax.set_ylim(0, 1.05)
    ax.set_ylabel("Score")
    ax.set_title("Per-class segmentation metrics")
    ax.grid(axis="y", alpha=0.25)
    ax.legend(loc="upper right")
    ax.bar_label(iou_bars, fmt="%.3f", fontsize=8, padding=2)
    ax.bar_label(dice_bars, fmt="%.3f", fontsize=8, padding=2)
    fig.tight_layout()
    fig.savefig(path_class_metrics, dpi=220, bbox_inches="tight")
    plt.close(fig)

    shares = [item["share"] for item in UNET_BENCHMARK["classes"]]
    colors = ["#2E74B5", "#D92D20", "#F79009"]
    fig, ax = plt.subplots(figsize=(8.8, 4.4))
    bars = ax.barh(classes, shares, color=colors)
    ax.set_xlim(0, 100)
    ax.set_xlabel("Share of ground-truth pixels (%)")
    ax.set_title("Ground-truth pixel support distribution")
    ax.grid(axis="x", alpha=0.25)
    ax.invert_yaxis()
    for bar, share, item in zip(bars, shares, UNET_BENCHMARK["classes"]):
        ax.text(
            min(share + 1.2, 96.0),
            bar.get_y() + bar.get_height() / 2,
            f"{share:.2f}% ({item['support']:,} px)",
            va="center",
            fontsize=9,
            color="#0B2545",
        )
    fig.tight_layout()
    fig.savefig(path_support, dpi=220, bbox_inches="tight")
    plt.close(fig)


def build_report():
    GENERATED_DIR.mkdir(parents=True, exist_ok=True)
    REPORT_DIR.mkdir(parents=True, exist_ok=True)

    active = read_json(ROOT / "benchmarks" / "reports" / "cleanliness_current_active_sam3_disabled_summary.json")
    candidate = read_json(ROOT / "benchmarks" / "reports" / "cleanliness_candidate_unet_sam3_disabled_summary.json")
    roboflow = read_json(ROOT / "benchmarks" / "reports" / "cleanliness_roboflow_merged_summary.json")

    system_png = GENERATED_DIR / "system_architecture.png"
    fusion_png = GENERATED_DIR / "model_fusion_architecture.png"
    retrain_png = GENERATED_DIR / "retrain_loop.png"
    rates_png = GENERATED_DIR / "benchmark_rates.png"
    latency_png = GENERATED_DIR / "benchmark_latency.png"
    unet_improvement_png = GENERATED_DIR / "unet_benchmark_improvement.png"
    unet_class_png = GENERATED_DIR / "unet_per_class_metrics.png"
    unet_support_png = GENERATED_DIR / "unet_class_support.png"
    create_system_diagram(system_png)
    create_fusion_diagram(fusion_png)
    create_retrain_diagram(retrain_png)
    create_metric_charts(rates_png, latency_png, active, candidate, roboflow)
    create_unet_benchmark_charts(unet_improvement_png, unet_class_png, unet_support_png)

    doc = Document()
    style_doc(doc)
    add_footer(doc)
    add_title_block(doc)

    doc.add_page_break()
    doc.add_heading("1. Executive Summary", level=1)
    doc.add_paragraph(
        "CleanOps AI Scoring is an operational computer-vision service for verifying post-cleaning image evidence. "
        "The service produces a cleanliness score, a PASS/PENDING/FAIL verdict, model evidence, and a visualization "
        "URL so supervisors can review uncertain or failed cases."
    )
    add_bullets(
        doc,
        [
            "The architecture is deliberately hybrid: YOLO handles object-level trash/debris; U-Net handles domain-specific dirty and wet-floor segmentation; SAM3/Roboflow-style segmentation provides broad auxiliary coverage.",
            "A refreshed U-Net pixel benchmark now evaluates 46/46 real images with ground-truth masks. The selected ResNet34 U-Net checkpoint reaches Pixel Accuracy 0.960080, Mean IoU 0.418637, and Mean Dice/F1 0.478676.",
            "The strongest U-Net improvement is on dirty_area: IoU 0.295541 and Dice/F1 0.456243. Wet_surface remains unsolved in this checkpoint and is called out as the next data priority.",
            "The service unions U-Net and auxiliary segmentation masks to reduce missed dirty cases, then applies explainable rule-based scoring and safety calibration.",
            "The current Roboflow/SAM3 auxiliary run achieved 0.00% false-pass on 18 evaluated pilot samples, but it routed every sample to PENDING review. This is a conservative safety result, not a claim of final autonomous accuracy.",
        ],
    )
    add_callout(
        doc,
        "Defense note: why this is not a single-model system",
        "Cleanliness evidence appears at different levels: loose objects, thin floor stains, wet reflections, and broad scene-level cues. "
        "A hybrid model stack is easier to validate, explain, and improve than a single black-box classifier.",
        fill=PALE_GOLD,
    )

    doc.add_heading("2. Service Purpose and Operational Output", level=1)
    doc.add_paragraph(
        "The service supports the CleanOps workflow after a cleaning task is completed. Instead of merely classifying an image as clean or dirty, it returns operational evidence that can be stored, audited, and reviewed by a supervisor."
    )
    add_table(
        doc,
        ["Output", "Meaning"],
        [
            ["quality_score", "Numerical cleanliness score derived from dirty/wet coverage and object penalty."],
            ["verdict", "Final operational decision: PASS, PENDING, or FAIL."],
            ["raw_verdict / calibration", "Explains when safety rules changed a raw model decision into a review-required case."],
            ["visualization.url", "Blob URL containing the overlay image used for demo, defense, and supervisor review."],
            ["sam3 block", "Compatibility contract for auxiliary segmentation, currently backed by Roboflow Workflow for lightweight runtime."],
        ],
        widths=[1.8, 4.5],
    )

    doc.add_heading("3. System Architecture", level=1)
    doc.add_paragraph(
        "The AI service is deployed as a FastAPI inference service behind the CleanOps backend. The backend owns business workflow and scoring jobs; the AI service owns inference, evidence fusion, scoring, and visualization generation."
    )
    doc.add_picture(str(system_png), width=Inches(6.25))
    add_caption(doc, "Figure 1. CleanOps system architecture: backend orchestration, AI scoring, blob visualization, review, retrain, and promotion gate.")
    add_table(
        doc,
        ["Layer", "Responsibility"],
        [
            ["CleanOps Backend", "Creates scoring jobs, stores results, exposes status and evidence to frontend/mobile clients."],
            ["AI Scoring API", "Loads images, runs YOLO/U-Net/auxiliary segmentation, computes score, and generates visualization."],
            ["Model and Blob Storage", "Stores active/candidate artifacts, runtime cache, and visualization images."],
            ["Human Review and Retrain", "Turns supervisor decisions and approved annotations into future candidate datasets."],
        ],
        widths=[2.0, 4.3],
    )

    doc.add_heading("4. AI Architecture and Model Fusion", level=1)
    doc.add_paragraph(
        "The AI stack separates evidence collection from business decision logic. This makes the system easier to explain during defense and safer to operate when one model is uncertain."
    )
    doc.add_picture(str(fusion_png), width=Inches(6.25))
    add_caption(doc, "Figure 2. Model fusion: YOLO, U-Net, and SAM3/Roboflow auxiliary segmentation are merged before scoring.")
    add_table(
        doc,
        ["Component", "Role", "Why it matters"],
        [
            ["YOLO", "Object-level detector for trash-like items and debris.", "Captures discrete evidence that segmentation can miss."],
            ["U-Net with ResNet encoder", "Domain-specific dirty/wet segmentation.", "Learns CleanOps floor patterns such as stains, residue, and wet surfaces."],
            ["SAM3/Roboflow-style auxiliary segmentation", "Foundation-style segmentation using classes such as Garbage, Stain, Stained_Floor, Wet_Floor.", "Adds general segmentation support when the specialized model under-detects."],
            ["Scoring rules", "Combines dirty coverage, wet coverage, object penalty, thresholds, and calibration.", "Keeps decisions explainable and auditable."],
        ],
        widths=[1.5, 2.2, 2.6],
    )

    doc.add_heading("5. Dataset and Training Provenance", level=1)
    doc.add_paragraph(
        "The repository separates raw/training data from frozen benchmark data. Frozen benchmark samples are kept under benchmarks and should not be used for training."
    )
    add_table(
        doc,
        ["Source", "Use in project", "Notes"],
        [
            ["Kaggle Trash Detection", "YOLO object detection data.", "Referenced in the repository downloader/docs."],
            ["Roboflow Clean/Unclean Floor", "Floor cleanliness images and object/cleanliness examples.", "Used as a public/raw data source, separate from the Roboflow Workflow provider."],
            ["Mendeley Indoor Waste Dataset", "Indoor waste/trash data.", "Manual download path documented in README."],
            ["Mendeley Stagnant Water Dataset", "Wet-floor/water examples.", "Useful for improving wet_surface annotation and segmentation."],
            ["Approved supervisor annotations", "Future production retrain source.", "Only approved labels should become retrain masks or YOLO boxes."],
        ],
        widths=[1.8, 2.0, 2.5],
    )
    doc.add_picture(str(retrain_png), width=Inches(6.25))
    add_caption(doc, "Figure 3. Retrain loop: review and approved annotations feed candidate training; promotion requires benchmark evidence.")

    doc.add_heading("6. Scoring Formula and Calibration", level=1)
    doc.add_paragraph("The core scoring formula is intentionally simple and explainable:")
    add_table(
        doc,
        ["Step", "Formula or rule"],
        [
            ["Mask fusion", "combined_dirty_coverage_pct is computed from the merged U-Net + auxiliary dirty/wet mask."],
            ["Base clean score", "base_clean_score = 100 - combined_dirty_coverage_pct"],
            ["Object penalty", "object_penalty = min(40, penalty_detections_count * 10)"],
            ["Quality score", "quality_score = clamp(base_clean_score - object_penalty, 0, 100)"],
            ["Verdict", "PASS above environment threshold; PENDING for review band; FAIL below review band."],
        ],
        widths=[1.8, 4.5],
    )
    add_callout(
        doc,
        "Important calibration interpretation",
        "When Roboflow/SAM3 produces an overly broad mask without strong supporting evidence, the service should route the case to PENDING review. "
        "This avoids false-pass while also avoiding an overconfident automatic FAIL.",
        fill=PALE_GOLD,
    )

    doc.add_heading("7. Benchmark Methodology and Results", level=1)
    doc.add_paragraph(
        "The reported metrics below come from real benchmark artifacts: the repository service-level reports and the refreshed 46-image U-Net benchmark report. No metrics are invented for this report."
    )

    doc.add_heading("7.1 Pixel-level U-Net Benchmark", level=2)
    doc.add_paragraph(
        "The refreshed pixel-level benchmark evaluates the selected U-Net checkpoint on the full "
        "data/processed/benchmark set: 46/46 real images with ground-truth masks and 0 missing masks. "
        "The checkpoint was selected by the project convention of choosing the run with the minimum "
        "best_val_loss parsed from training logs."
    )
    add_table(
        doc,
        ["Field", "Value"],
        [
            ["Benchmark dataset", UNET_BENCHMARK["dataset"]],
            ["Images evaluated", str(UNET_BENCHMARK["evaluated"])],
            ["Missing masks", str(UNET_BENCHMARK["missing_masks"])],
            ["Selected checkpoint", UNET_BENCHMARK["checkpoint"]],
            ["Selection rule", UNET_BENCHMARK["selection_rule"]],
            ["Best validation loss", f"{UNET_BENCHMARK['best_val_loss']:.4f}"],
            ["Source report", str(UNET_BENCHMARK_SOURCE)],
        ],
        widths=[1.7, 4.6],
    )
    add_table(
        doc,
        ["Metric", "Previous benchmark", "New benchmark", "Change"],
        [
            ["Pixel Accuracy", "0.947288", "0.960080", "+0.012792"],
            ["Mean IoU", "0.372956", "0.418637", "+0.045681"],
            ["Mean Dice/F1", "0.422034", "0.478676", "+0.056642"],
        ],
        widths=[2.1, 1.3, 1.3, 1.3],
    )
    doc.add_picture(str(unet_improvement_png), width=Inches(6.0))
    add_caption(doc, "Figure 4. U-Net pixel benchmark improvement using the best-val-loss checkpoint on the 46-image mask benchmark.")
    doc.add_paragraph(
        "Pixel Accuracy is high partly because background dominates the pixel distribution. For this task, Mean IoU and Dice/F1 are more informative because the important warning classes are small foreground regions."
    )
    add_table(
        doc,
        ["Class", "Support pixels", "Pixel share", "IoU", "Dice/F1"],
        [
            ["background", "604,257,149", "95.28%", "0.960370", "0.979785"],
            ["dirty_area", "15,416,922", "2.43%", "0.295541", "0.456243"],
            ["wet_surface", "14,487,258", "2.28%", "0.000000", "0.000000"],
        ],
        widths=[1.4, 1.25, 1.25, 1.3, 1.3],
    )
    doc.add_picture(str(unet_class_png), width=Inches(6.0))
    add_caption(doc, "Figure 5. Per-class U-Net metrics. The checkpoint learns dirty_area but still fails to segment wet_surface.")
    doc.add_picture(str(unet_support_png), width=Inches(5.5))
    add_caption(doc, "Figure 6. Ground-truth class imbalance across 634,161,329 pixels.")
    add_table(
        doc,
        ["Ground truth / Prediction", "background", "dirty_area", "wet_surface"],
        [
            ["background", "604,141,410", "115,739", "0"],
            ["dirty_area", "10,712,425", "4,704,497", "0"],
            ["wet_surface", "14,101,657", "385,601", "0"],
        ],
        widths=[2.0, 1.35, 1.35, 1.35],
    )
    add_callout(
        doc,
        "Interpretation of the refreshed U-Net benchmark",
        "The new checkpoint provides real evidence of learning dirty_area. However, it still predicts no wet_surface pixels, so wet-floor detection should remain a data and training priority rather than a claimed strength.",
        fill=PALE_GOLD,
    )

    doc.add_heading("7.2 Pilot Business Benchmark", level=2)
    doc.add_paragraph(
        "The service-level benchmark below is separate from the U-Net pixel benchmark. It measures operational PASS/PENDING/FAIL behavior through the live API, including scoring rules, calibration, visualization generation, and external Roboflow/SAM3-style auxiliary segmentation. It should not be read as a pixel-level mIoU benchmark for Roboflow/SAM3 because no frozen auxiliary mask ground truth is available."
    )
    benchmark_rows = [
        [
            "Current active, SAM3 disabled",
            str(active["evaluated_samples"]),
            str(active["skipped_samples"]),
            pct(active["verdict_accuracy"]),
            pct(active["false_pass_rate"]),
            pct(active["false_fail_rate"]),
            pct(active["pending_review_rate"]),
            f"{active['average_latency_ms']:.2f} ms",
        ],
        [
            "Candidate U-Net, SAM3 disabled",
            str(candidate["evaluated_samples"]),
            str(candidate["skipped_samples"]),
            pct(candidate["verdict_accuracy"]),
            pct(candidate["false_pass_rate"]),
            pct(candidate["false_fail_rate"]),
            pct(candidate["pending_review_rate"]),
            f"{candidate['average_latency_ms']:.2f} ms",
        ],
        [
            "Roboflow/SAM3 merged + calibration",
            str(roboflow["evaluated_samples"]),
            str(roboflow["skipped_samples"]),
            pct(roboflow["verdict_accuracy"]),
            pct(roboflow["false_pass_rate"]),
            pct(roboflow["false_fail_rate"]),
            pct(roboflow["pending_review_rate"]),
            f"{roboflow['average_latency_ms']:.2f} ms",
        ],
    ]
    add_table(
        doc,
        ["Variant", "Eval.", "Skip", "Accuracy", "False pass", "False fail", "Pending", "Avg latency"],
        benchmark_rows,
        widths=[2.0, 0.55, 0.55, 0.75, 0.8, 0.8, 0.75, 0.9],
    )
    doc.add_picture(str(rates_png), width=Inches(6.25))
    add_caption(doc, "Figure 7. Pilot benchmark rates by variant. The Roboflow merged run eliminates false-pass in this small pilot but sends all evaluated cases to review.")
    doc.add_picture(str(latency_png), width=Inches(6.0))
    add_caption(doc, "Figure 8. Average API latency by variant. Roboflow adds external workflow latency compared with U-Net-only variants.")

    doc.add_heading("7.3 Roboflow/SAM3 Merged Run", level=2)
    cm = roboflow["confusion_matrix"]
    add_table(
        doc,
        ["Expected / Predicted", "PASS", "PENDING", "FAIL"],
        [
            ["PASS", str(cm["PASS"]["PASS"]), str(cm["PASS"]["PENDING"]), str(cm["PASS"]["FAIL"])],
            ["PENDING", str(cm["PENDING"]["PASS"]), str(cm["PENDING"]["PENDING"]), str(cm["PENDING"]["FAIL"])],
            ["FAIL", str(cm["FAIL"]["PASS"]), str(cm["FAIL"]["PENDING"]), str(cm["FAIL"]["FAIL"])],
        ],
        widths=[2.0, 1.0, 1.0, 1.0],
    )
    add_table(
        doc,
        ["Metric", "Value"],
        [
            ["Input rows", str(roboflow["input_rows"])],
            ["Evaluated samples", str(roboflow["evaluated_samples"])],
            ["Skipped samples", str(roboflow["skipped_samples"]) + " (403 image URL)"],
            ["False-pass rate", pct(roboflow["false_pass_rate"])],
            ["False-fail rate", pct(roboflow["false_fail_rate"])],
            ["Pending review rate", pct(roboflow["pending_review_rate"])],
            ["Calibrated rate", pct(roboflow["calibrated_rate"])],
            ["Average SAM3/Roboflow elapsed", f"{roboflow['average_sam3_elapsed_ms']:.2f} ms"],
            ["SAM3/Roboflow status counts", str(roboflow["sam3_status_counts"])],
            ["Dirty coverage source counts", str(roboflow["dirty_coverage_source_counts"])],
        ],
        widths=[2.6, 3.7],
    )

    doc.add_heading("8. Interpretation and Promotion Decision", level=1)
    add_bullets(
        doc,
        [
            "The candidate U-Net should not be promoted because false-pass increased from 11.11% to 38.89% on the pilot benchmark.",
            "The Roboflow/SAM3 merged run reduces false-pass to 0.00% in the pilot, but its pending review rate is 100.00%, so it should be presented as a safety-oriented auxiliary integration rather than a finished autonomous classifier.",
            "The refreshed 46-image U-Net benchmark strengthens the technical evidence for dirty_area segmentation, but the service still needs more wet_surface data before wet-floor segmentation can be defended as reliable.",
            "Broad auxiliary masks are useful as warning evidence, but they require calibration and supervisor review to avoid overconfident FAIL decisions.",
            "The most defensible current position is: the integration is working, the architecture is explainable, and further benchmark/annotation work is required before production promotion.",
        ],
    )
    add_callout(
        doc,
        "Defense note: why rule-based scoring is used",
        "A black-box binary answer is hard to defend when a supervisor asks why a job failed. Rule-based scoring keeps each penalty traceable: dirty coverage, wet coverage, object penalty, environment threshold, and calibration reason.",
        fill=PALE_BLUE,
    )

    doc.add_heading("9. Known Limitations", level=1)
    add_bullets(
        doc,
        [
            "The pilot benchmark is small: 19 rows, 18 evaluated, and 1 skipped because the source image returned HTTP 403.",
            "The U-Net mask benchmark is stronger than before at 46 images, but it is still class-imbalanced: background accounts for 95.28% of ground-truth pixels.",
            "There is no frozen golden mask/polygon benchmark for SAM3/Roboflow yet, so the auxiliary provider is evaluated by business verdicts rather than pixel-level mIoU.",
            "The wet_surface class remains unsolved in the selected U-Net checkpoint and needs more approved annotations across different floor materials and lighting conditions.",
            "SAM3 local GPU was blocked by a CUDA requirement mismatch: the container required CUDA >= 12.8 while the local driver reported CUDA 12.3.",
            "Roboflow is the practical lightweight provider for demo/runtime, but external workflow latency and broad masks must be monitored.",
        ],
    )

    doc.add_heading("10. Recommended Next Steps", level=1)
    add_bullets(
        doc,
        [
            "Implement the score display consistency fix so quality score remains the raw model score while calibration only changes verdict and review metadata.",
            "Add more wet-floor examples and reviewed masks, then retrain with class-weighted Dice/Focal loss to reduce the influence of background dominance.",
            "Re-run the same 46-image U-Net benchmark after each fine-tune so checkpoint comparisons remain fair and repeatable.",
            "Expand the service-level pilot benchmark to 30-50 real images with balanced clean, slightly dirty, and obviously dirty cases.",
            "Use SAM3/Roboflow to accelerate mask proposals, but require human review before those masks become training labels.",
            "Tune Roboflow class prompts and confidence thresholds to reduce broad masks before lowering the pending review rate.",
            "Promote only candidates that pass the gate: false-pass must not increase, verdict accuracy must not decrease, and runtime latency must remain acceptable.",
        ],
    )

    doc.add_heading("11. Evidence and Source Artifacts", level=1)
    add_table(
        doc,
        ["Artifact", "Purpose"],
        [
            ["benchmarks/reports/cleanliness_ab_comparison.md", "A/B comparison between active and candidate U-Net."],
            ["benchmarks/reports/cleanliness_current_active_sam3_disabled_summary.json", "Baseline U-Net-only business metrics."],
            ["benchmarks/reports/cleanliness_candidate_unet_sam3_disabled_summary.json", "Candidate U-Net business metrics."],
            ["benchmarks/reports/cleanliness_roboflow_merged_summary.json", "Roboflow/SAM3 auxiliary merged benchmark metrics."],
            ["benchmarks/reports/sam3_smoke_blocker_20260527.md", "Local SAM3 CUDA runtime blocker evidence."],
            [str(UNET_BENCHMARK_SOURCE), "New 46-image U-Net mask benchmark source report."],
            ["docs/report/benchmark_report_vi.md", "Vietnamese source report and benchmark interpretation."],
        ],
        widths=[3.2, 3.1],
    )

    doc.save(OUT_DOCX)
    print(OUT_DOCX)


if __name__ == "__main__":
    build_report()
